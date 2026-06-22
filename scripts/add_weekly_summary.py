#!/usr/bin/env python3
"""为 recording 目录下的常规会议记录添加/更新「上周总结」行。"""

from __future__ import annotations

import re
import shutil
import sys
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import Table

BASE_DIR = Path(__file__).resolve().parent.parent
RECORDING_DIR = BASE_DIR / "recording"

# 会议序号 -> docx 相对路径
MEETING_FILES: dict[int, str] = {
    1: "3_03第一次会议/会议记录.docx",
    2: "3_10第二次会议/会议记录.docx",
    3: "3_17第三次会议/会议记录.docx",
    4: "3_24第四次会议/会议记录.docx",
    5: "3_31第五次会议/会议记录.docx",
    6: "4_07第六次会议/会议记录.docx",
    7: "4_14第七次会议/会议记录.docx",
    8: "4_21第八次会议/会议记录.docx",
    9: "4_28第九次会议/第九次会议纪要.docx",
    10: "5_03第十次会议/第十次会议纪要.docx",
    11: "5_05第十一次会议/第十一次会议纪要.docx",
    12: "5_12第十二次会议/会议纪要.docx",
    13: "5_19第十三次会议/第十三次会议纪要.docx",
    14: "5_26第十四次会议/会议纪要.docx",
    15: "5_29第十五次会议/会议纪要.docx",
    16: "6_06第十六次会议/第十六次会议会议纪要.docx",
    17: "6_09第十七次会议/第十七次会议纪要.docx",
    18: "6_21第十八次会议/第十八次会议纪要.docx",
}

WEEKLY_SUMMARIES: dict[int, str] = {
    1: "无（首次会议）",
    2: (
        "上次启动会议已完成竞品调研与产品方向讨论：确定以华为云 ModelArts 为核心对标，"
        "形成大众层、专业层、多智能体三层产品架构思路。"
        "会后布置的客户群识别、项目愿景文件研究、SRS 文件研究、需求跟踪矩阵配置、"
        "需求项目计划与甘特图制定等任务尚处于启动阶段，尚未形成正式文档产出。"
    ),
    3: (
        "上次布置的 8 项文档初稿任务推进情况如下：项目计划（陈鹏宇）、项目章程（李俊翰/张明慧）、"
        "需求分析文档（陈鹏宇）等核心文档已启动编写；任务清单上传 GitHub（曹志隆）、"
        "PPT 框架整理（陈天阳）同步进行中。"
        "整体采用 AI 生成初稿 + 人工调整策略，但 GitHub 协作规范与文档格式尚未完全统一，"
        "多项文档仍待周三晚前提交初稿，可行性分析、需求文档格式等细节待与老师确认。"
    ),
    4: (
        "上次 6 项任务部分完成：GitHub 分支已按人名缩写创建（曹志隆）；"
        "PPT 统一模板已寻找（李俊翰）；WBS 思维导图已有初步进展（张明慧）；"
        "UML 概述内容范围已向老师确认（陈鹏宇）。"
        "甘特图（含人力资源分配）与 PPT 内容填充仍在推进，"
        "翻转课堂 PPT 制作（陈鹏宇/陈天阳/梅怡钿）整体进度待小组会议进一步对齐。"
    ),
    5: (
        "上次翻转课堂任务已通过抽签接龙完成分工：张明慧负责 PPT 展示与内容大纲，"
        "陈天阳与曹志隆负责 PPT 制作，梅怡钿与陈鹏宇负责需求计划内容修订，"
        "李俊翰负责 PPT 优化与组合。主讲人已定，大纲框架已确定，"
        "PPT 总页数控制在 40 页左右，各模块内容填充与文档修订工作已启动。"
    ),
    6: (
        "上次计划书修订任务部分完成：沟通管理、成本（预算）管理、范围管理三类子计划"
        "确认可直接从原计划书复用提取，采购管理计划模块已一致同意删除。"
        "干系人需求调研已启动，要求每位成员寻找 1–2 名潜在用户进行调研；"
        "「获取需求」专项文档编撰（梅怡钿/陈鹏宇）与 PPT 讲稿优化（张明慧/陈天阳）已分配。"
        "部分 Word 文档内容组员已初步完成但尚未全部提交至群。"
    ),
    7: (
        "上次文档与甘特图任务部分完成：各组员已初步完成部分 Word 文档内容，"
        "但尚未全部提交至群，仍需修改完善；甘特图已出初稿，经老师指导后需进一步调整。"
        "李俊翰已指定负责甘特图修改，PPT 制作（陈鹏宇/陈天阳/梅怡钿）与演讲人（陈天阳）分工已明确，"
        "本周优先处理未完成的上周任务，暂不推进后续新内容。"
    ),
    8: (
        "上次需求获取阶段任务产出已汇总至群内：梅怡钿已完成用户群分类；"
        "愿景文档、用例、用户工作流程、质量属性补充、开发原型及关联图等任务"
        "通过抽签分配给全体成员，各项产出已有初步成果。"
        "陈天阳等组员已完成部分提交，但部分文档内容仍需修正完善，"
        "尚未全部达到可评审标准。"
    ),
    9: (
        "上次任务持续推进中：用户邀请函已由曹志隆制作并以文字形式发给第十组（用户代表）；"
        "顶层用例图修改与 SRS 初稿撰写仍在进行；梅怡钿负责的个人作业（邀请函、用例图等）已启动。"
        "经讨论团队已收窄系统定位，明确先做可演示的简单原型而非大模型复现，"
        "原型设计方向基本明确，API 调用方等不合理设计待删除调整。"
    ),
    10: (
        "上次任务取得阶段性成果：曹志隆与张明慧已完成基本原型，界面设计获团队一致认可并暂定当前版本；"
        "项目愿景与范围文档修订（陈鹏宇）、项目发起人邀请函制作（陈鹏宇）持续推进；"
        "用户代表（第十组、第二组）需求确认同步事宜已安排；"
        "问卷数据收集与分析（梅怡钿）工作已启动。"
    ),
    11: (
        "上次任务进展如下：问卷已发给其他小组但填写人数不足（仅约 3 人），"
        "需加强督促；向第二组、第十组发送的用户邀请函草稿仍需调整，尚未正式发出；"
        "项目发起人邀请函（发给老师）与用户邀请函（发给用户代表）的区别已澄清。"
        "RUP 方法学习与会议记录整理（李俊翰）、按新用户分类修改愿景文档（李俊翰/陈鹏宇）仍在进行中。"
    ),
    12: (
        "上次《项目视图与范围》文档按评审意见进行结构性修订：业务风险、第三方服务、"
        "主要特性、发布范围（三版本规划）、部署考虑等核心章节已讨论修改方案；"
        "主要特性/发布范围难点讨论（李俊翰/陈鹏宇/曹志隆）、问卷数据扩展（陈天阳）、"
        "核心用例撰写（陈天阳/梅怡钿/张明慧）及与老师访谈邀请（陈鹏宇）均已安排推进。"
    ),
    13: (
        "上次功能范围裁剪与业务流程梳理已完成：遵照老师意见剔除「运维部署」模块，"
        "确定平台产出物为可下载模型文件，简化模型训练管理并与模型开发合并；"
        "普通用户与管理员两条核心业务流程已梳理清楚。"
        "SRS、用例图、状态图更新与原型设计已按抽签分工启动，"
        "管理员模块功能细化（李俊翰/曹志隆）及与第二组、第十组的功能优先级调研对接已安排。"
    ),
    14: (
        "上次苏老师新需求已响应并完成分工：陈鹏宇负责修改范围文档、明确团队协作功能限制；"
        "李俊翰负责修改原型（增加镜像管理、共享导入等功能）及 CSCI 部分详细分工；"
        "陈天阳负责翻转课堂大纲与演讲内容，张明慧负责 SRS 评审 PPT。"
        "镜像管理（保存为镜像/从镜像恢复）、共享导入与文件导入等新功能点已纳入 SRS 调整范围，"
        "CSCI 功能需求部分需全面返工。"
    ),
    15: (
        "上次导师评审意见对齐后，多项任务并行推进：曹志隆/陈鹏宇联系苏老师与第十组敲定 JAD 会议；"
        "陈天阳优化 QFD 问卷，曹志隆填写管理员调查问卷；曹志隆将 SRS 按用户/管理员维度拆分，"
        "陈鹏宇更新顶层用例图与对话框图，梅怡钿重构测试用例，陈天阳编写用户手册，"
        "李俊翰汇总答辩 PPT 框架划分。"
        "JAD 会议、问卷反馈收集与功能优先级合分等待 JAD 后进一步完成。"
    ),
    16: (
        "上次 SRS 功能点细化与非功能性需求撰写持续推进：「管理数据」等功能已逐项拆分，"
        "大文件上传/下载、训练监控等可行性已讨论确认；李俊翰负责更新功能点列表、"
        "撰写非功能性需求及补充 SRS 缺失章节，陈鹏宇对接第十组补充需求确认文件。"
        "SRS 功能需求整体具备可行性，非功能性需求（性能、并发、可靠性）已确认需新增。"
    ),
    17: (
        "上次 SRS 文档完善工作取得较大进展：SRS 已达 146 页，测试用例文档已从 SRS 中独立拆分；"
        "课程大作业评审检查表 30 项已逐项核对，Vision & Scope、Context Diagram、"
        "用户群分类、JAD 会议记录、用例文档、非功能性需求等大部分条目已确认完成。"
        "对话框图补充（曹志隆）、可行性分析合并（李俊翰）、用户手册红框标注（陈天阳）、"
        "需求优先级打分（张明慧）等待办已分配，部分插图完整性仍待检查。"
    ),
    18: (
        "上次 SRS 定稿与需求变更工作全面推进：SRS 170 页定稿邀请已发起，全员需完成确认以确立需求基线；"
        "新增用户个人信息优化需求（头像上传/修改、昵称替代实名、自定义显示名称）已纳入变更流程，"
        "判定为功能新增且与现有基线无冲突。"
        "CCB 会议已安排（周五晚 9:30），配套变更申请报告、可行性分析、原型更新、"
        "测试用例迭代、用户手册修订、飞书需求管理工具完善等任务已全部分工，"
        "大部分任务要求本周末前完成。"
    ),
}

SUMMARY_LABELS = {"上周总结", "上一周的总结", "会议上一周的总结"}


def cell_text(row_cells: list, index: int = 0) -> str:
    if index >= len(row_cells):
        return ""
    return row_cells[index].text.strip()


def set_cell_text_preserve_style(cell, text: str) -> None:
    """仅替换单元格文本，尽量保留段落/字体样式。"""
    if not cell.paragraphs:
        cell.text = text
        return
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.text = text
    for extra in cell.paragraphs[1:]:
        extra.clear()


def find_meeting_table(doc: Document) -> Table | None:
    for table in doc.tables:
        if table.rows and cell_text(table.rows[0].cells, 0) == "会议名称":
            return table
    return None


def find_row_index(table: Table, label: str) -> int | None:
    for i, row in enumerate(table.rows):
        if cell_text(row.cells, 0) == label:
            return i
    return None


def find_attendees_row_index(table: Table) -> int:
    idx = find_row_index(table, "与会人员")
    if idx is not None:
        return idx
    raise ValueError("未找到「与会人员」行，无法复制行格式")


def insert_row_after(table: Table, after_index: int, template_index: int) -> None:
    template_tr = table.rows[template_index]._tr
    new_tr = deepcopy(template_tr)
    table.rows[after_index]._tr.addnext(new_tr)


def upsert_weekly_summary(table: Table, summary_text: str) -> str:
    """返回操作类型：updated / inserted"""
    theme_idx = find_row_index(table, "会议主题")
    if theme_idx is None:
        raise ValueError("未找到「会议主题」行")

    next_idx = theme_idx + 1
    if next_idx < len(table.rows):
        next_label = cell_text(table.rows[next_idx].cells, 0)
        if next_label in SUMMARY_LABELS:
            row = table.rows[next_idx]
            set_cell_text_preserve_style(row.cells[0], "上周总结")
            set_cell_text_preserve_style(row.cells[1], summary_text)
            return "updated"

    attendees_idx = find_attendees_row_index(table)
    insert_row_after(table, theme_idx, attendees_idx)
    new_row = table.rows[theme_idx + 1]
    set_cell_text_preserve_style(new_row.cells[0], "上周总结")
    set_cell_text_preserve_style(new_row.cells[1], summary_text)
    return "inserted"


def process_meeting(meeting_num: int, rel_path: str) -> dict:
    path = RECORDING_DIR / rel_path
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    summary = WEEKLY_SUMMARIES[meeting_num]
    doc = Document(str(path))
    table = find_meeting_table(doc)
    if table is None:
        raise ValueError(f"未找到会议信息表格: {path}")

    action = upsert_weekly_summary(table, summary)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir=path.parent) as tmp:
        tmp_path = Path(tmp.name)
    try:
        doc.save(str(tmp_path))
        try:
            shutil.move(str(tmp_path), str(path))
            saved_as = str(path)
        except PermissionError:
            fallback = path.with_name(f"{path.stem}_updated{path.suffix}")
            if fallback.exists():
                fallback.unlink()
            shutil.move(str(tmp_path), str(fallback))
            saved_as = str(fallback)
            action = f"{action}_fallback"
    except Exception:
        if tmp_path.exists():
            tmp_path.unlink()
        raise

    return {
        "meeting": meeting_num,
        "path": rel_path,
        "action": action,
        "summary": summary,
        "saved_as": saved_as,
    }


def main() -> int:
    results = []
    errors = []

    for num in sorted(MEETING_FILES.keys()):
        rel = MEETING_FILES[num]
        try:
            result = process_meeting(num, rel)
            results.append(result)
            saved = result.get("saved_as", rel)
            print(f"[OK] 第{num:02d}次 ({result['action']}): {saved}")
        except Exception as exc:  # noqa: BLE001
            errors.append((num, rel, str(exc)))
            print(f"[ERR] 第{num:02d}次: {rel} -> {exc}", file=sys.stderr)

    print("\n--- 处理结果对照 ---")
    for r in results:
        preview = r["summary"][:50] + ("..." if len(r["summary"]) > 50 else "")
        print(f"第{r['meeting']:02d}次 [{r['action']}]: {preview}")

    if errors:
        print(f"\n失败 {len(errors)} 个文件", file=sys.stderr)
        return 1
    print(f"\n成功处理 {len(results)} 个文件")
    return 0


if __name__ == "__main__":
    sys.exit(main())
